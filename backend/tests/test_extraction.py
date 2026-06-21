import io

from app.extraction import extract_pages, extract_text


class TestExtractText:
    def test_extracts_plain_text(self):
        content = b"Hello world, this is plain text."
        assert extract_text(content, "text/plain") == "Hello world, this is plain text."

    def test_extracts_markdown_as_plain_text(self):
        content = b"# Heading\n\nSome **bold** text."
        assert extract_text(content, "text/markdown") == "# Heading\n\nSome **bold** text."

    def test_extracts_text_from_pdf(self):
        from pypdf import PdfWriter

        writer = PdfWriter()
        page = writer.add_blank_page(width=200, height=200)
        page.merge_page(_pdf_page_with_text("Hello from PDF"))
        buf = io.BytesIO()
        writer.write(buf)
        content = buf.getvalue()

        text = extract_text(content, "application/pdf")
        assert "Hello from PDF" in text

    def test_extracts_text_from_docx(self):
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Hello from DOCX")
        buf = io.BytesIO()
        doc.save(buf)
        content = buf.getvalue()

        text = extract_text(content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        assert "Hello from DOCX" in text

    def test_raises_for_unsupported_content_type(self):
        import pytest

        with pytest.raises(ValueError):
            extract_text(b"data", "application/x-msdownload")


class TestExtractPages:
    def test_plain_text_is_a_single_page_with_no_page_number(self):
        content = b"Hello world."
        assert extract_pages(content, "text/plain") == [(None, "Hello world.")]

    def test_docx_is_a_single_page_with_no_page_number(self):
        from docx import Document as DocxDocument

        doc = DocxDocument()
        doc.add_paragraph("Hello from DOCX")
        buf = io.BytesIO()
        doc.save(buf)

        pages = extract_pages(
            buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        assert len(pages) == 1
        page_number, text = pages[0]
        assert page_number is None
        assert "Hello from DOCX" in text

    def test_pdf_returns_one_entry_per_page_numbered_from_one(self):
        from pypdf import PdfWriter

        writer = PdfWriter()
        for label in ("First page", "Second page"):
            page = writer.add_blank_page(width=200, height=200)
            page.merge_page(_pdf_page_with_text(label))
        buf = io.BytesIO()
        writer.write(buf)

        pages = extract_pages(buf.getvalue(), "application/pdf")

        assert [p[0] for p in pages] == [1, 2]
        assert "First page" in pages[0][1]
        assert "Second page" in pages[1][1]

    def test_raises_for_unsupported_content_type(self):
        import pytest

        with pytest.raises(ValueError):
            extract_pages(b"data", "application/x-msdownload")


def _pdf_page_with_text(text: str):
    from reportlab.pdfgen import canvas
    from pypdf import PdfReader

    buf = io.BytesIO()
    c = canvas.Canvas(buf)
    c.drawString(50, 100, text)
    c.save()
    buf.seek(0)
    return PdfReader(buf).pages[0]
